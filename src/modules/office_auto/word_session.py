# word_session.py

import os
import re
import comtypes.client
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WordSession:
    def __init__(self):
        self.doc = None
        self.is_active = False
        self.download_path = os.path.join(os.path.expanduser("~"), "Downloads")

    def _speak_local(self, text):
        try:
            import pyttsx3, pythoncom
            pythoncom.CoInitialize()
            engine = pyttsx3.init("sapi5")
            engine.setProperty('rate', 150)
            voices = engine.getProperty('voices')
            for v in voices:
                if "spanish" in v.name.lower() or "español" in v.name.lower():
                    engine.setProperty('voice', v.id)
                    break
            engine.say(text)
            engine.runAndWait()
        except: pass
        finally:
            try: pythoncom.CoUninitialize()
            except: pass

    def start_session(self, initial_filename: str = None) -> str:
        self.doc = Document()
        self.is_active = True
        self._speak_local("Editor listo. Di 'Título', 'Párrafo', 'Oración' o 'Dictado'.")

    def _normalize_punctuation(self, text: str) -> str:
        text = text.replace("\r", "").replace("\n", " ").strip()
        rules = [
            (r"\s+\bpunto y coma\b", ";"), (r"\s+\bdos puntos\b", ":"),
            (r"\s+\bcoma\b", ","), (r"\s+\bpunto\b", "."),
            (r"\s+\babre interrogación\b", " ¿"), (r"\s+\bcierra interrogación\b", "?"),
            (r"\s+\babre paréntesis\b", " ("), (r"\s+\bcierra paréntesis\b", ")"),
            (r"\s+\bcomillas\b", '"')
        ]
        for pattern, replacement in rules:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        text = text.strip()
        if text: text = text[0].upper() + text[1:]
        return text

    def _add_styled_paragraph(self, text, style_type):
        if not self.doc or not text: return
        
        clean_text = self._normalize_punctuation(text)
        
        if style_type in ["parrafo", "oracion"]:
            if not clean_text.endswith((".", "?", "!", ":", ";", '"')):
                clean_text += "."

        if style_type == "oracion":
            if len(self.doc.paragraphs) > 0:
                p = self.doc.paragraphs[-1]
                prev_text = p.text.strip()
                if prev_text and not prev_text.endswith(" "):
                     p.add_run(" ")
                p.add_run(clean_text)
            else:
                self._create_paragraph(clean_text, 12, False, "JUSTIFY")
            return

        if style_type == "titulo": self._create_paragraph(clean_text, 16, True, "CENTER")
        elif style_type == "subtitulo": self._create_paragraph(clean_text, 14, True, "LEFT")
        else: self._create_paragraph(clean_text, 12, False, "JUSTIFY")

    def _create_paragraph(self, text, size, bold, align):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if align == "CENTER": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "LEFT": p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == "JUSTIFY": p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _convert_to_pdf(self, docx_path, pdf_path):
        try:
            import comtypes.gen
            word_app = comtypes.client.CreateObject('Word.Application')
            word_app.Visible = False
            doc = word_app.Documents.Open(docx_path)
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close()
            word_app.Quit()
            return True
        except: return False

    def _save_file(self, raw_input):
        format_type = "pdf" if "pdf" in raw_input.lower() else "docx"
        
        temp_name = raw_input.lower()
        remove_list = ["guardar documento", "guardar archivo", "guardar el documento",
                       "con el nombre", "llamado", "titulado", "como",
                       "en pdf", "formato pdf", "pdf", "en word"]
        
        for phrase in remove_list:
            temp_name = temp_name.replace(phrase, "")
        
        clean_name = re.sub(r'[<>:"/\\|?*]', '', temp_name).strip().capitalize()
        if not clean_name: clean_name = "Documento_Generado"
        
        docx_path = os.path.join(self.download_path, f"{clean_name}.docx")
        
        try:
            self.doc.save(docx_path)
            msg = f"Guardado {clean_name}."
            
            if format_type == "pdf":
                pdf_path = os.path.join(self.download_path, f"{clean_name}.pdf")
                success = self._convert_to_pdf(docx_path, pdf_path)
                if success: msg += " PDF creado."
                else: msg += " Solo DOCX."

            self.is_active = False
            self.doc = None
            self._speak_local(msg)

        except Exception as e:
            self._speak_local("Error al guardar.")

    def process_dictation(self, text: str):
        if not self.is_active: return

        text_lower = text.lower().strip()

        if "guardar" in text_lower:
            return self._save_file(text_lower)

        parts = text.split(' ', 1)
        if len(parts) < 2: return 

        content = parts[1]
        
        if text_lower.startswith(("título", "titulo")):
            self._add_styled_paragraph(content, "titulo")
            self._speak_local("Título.")

        elif text_lower.startswith(("subtítulo", "subtitulo")):
            self._add_styled_paragraph(content, "subtitulo")
            self._speak_local("Subtítulo.")

        elif text_lower.startswith(("párrafo", "parrafo")):
            self._add_styled_paragraph(content, "parrafo")
            self._speak_local("Párrafo.")
            
        elif text_lower.startswith(("oración", "oracion", "frase")):
            self._add_styled_paragraph(content, "oracion")
            self._speak_local("Punto seguido.")

        elif text_lower.startswith(("dictado", "escribir")):
            self._add_styled_paragraph(content, "parrafo")
            self._speak_local("Escrito.")

        else:
            print(f"   [WORD IGNORADO] '{text}'")
            return 

word_session = WordSession()